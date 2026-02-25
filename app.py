import customtkinter as ctk
import tkinter.messagebox
import tkinter.filedialog
from docx import Document
import os
from datetime import datetime

# Set the appearance and color theme
ctk.set_appearance_mode("System")  # Modes: "System" (standard), "Dark", "Light"
ctk.set_default_color_theme("blue")  # Themes: "blue" (standard), "green", "dark-blue"

class App(ctk.CTk):
    def __init__(self):
        super().__init__()

        # Window Configuration
        self.title("Document Generator Pro")
        self.geometry(f"{600}x{500}")
        self.resizable(False, False)
        self.last_saved_dir = os.path.expanduser("~/Documents") # Ensure a safe default path

        # ------------------ UI Layout ------------------

        # Header Label
        self.header_label = ctk.CTkLabel(self, text="Create a New Word Document", font=ctk.CTkFont(size=24, weight="bold"))
        self.header_label.pack(pady=(20, 10))

        # Title Input
        self.title_frame = ctk.CTkFrame(self)
        self.title_frame.pack(pady=10, padx=20, fill="x")
        
        self.doc_title_label = ctk.CTkLabel(self.title_frame, text="Document Title:")
        self.doc_title_label.pack(side="left", padx=10)
        
        self.doc_title_entry = ctk.CTkEntry(self.title_frame, placeholder_text="Enter the main heading...", width=300)
        self.doc_title_entry.pack(side="left", padx=10, fill="x", expand=True)

        # Content Input
        self.content_frame = ctk.CTkFrame(self)
        self.content_frame.pack(pady=10, padx=20, fill="both", expand=True)
        
        self.content_label = ctk.CTkLabel(self.content_frame, text="Document Body Content:")
        self.content_label.pack(anchor="w", padx=10, pady=(10, 0))
        
        self.content_textbox = ctk.CTkTextbox(self.content_frame, height=150)
        self.content_textbox.pack(padx=10, pady=10, fill="both", expand=True)

        # File Name Input (Now just suggests the name)
        self.filename_frame = ctk.CTkFrame(self)
        self.filename_frame.pack(pady=10, padx=20, fill="x")
        
        self.filename_label = ctk.CTkLabel(self.filename_frame, text="Suggested File Name:")
        self.filename_label.pack(side="left", padx=10)
        
        self.filename_entry = ctk.CTkEntry(self.filename_frame, placeholder_text="my_document", width=250)
        self.filename_entry.pack(side="left", padx=10, fill="x", expand=True)
        self.filename_entry.insert(0, f"Generated_Doc_{datetime.now().strftime('%Y%m%d%H%M%S')}")

        # Action Buttons
        self.button_frame = ctk.CTkFrame(self, fg_color="transparent")
        self.button_frame.pack(pady=(10, 20), padx=20, fill="x")

        self.generate_btn = ctk.CTkButton(
            self.button_frame, 
            text="✨ Generate Document ✨", 
            font=ctk.CTkFont(size=15, weight="bold"),
            height=40,
            command=self.generate_document
        )
        self.generate_btn.pack(side="left", expand=True, fill="x", padx=10)

        self.open_btn = ctk.CTkButton(
            self.button_frame, 
            text="📂 Open Save Folder", 
            fg_color="gray",
            hover_color="darkgray",
            height=40,
            command=self.open_folder
        )
        self.open_btn.pack(side="right", fill="x", padx=10)

    # ------------------ Action Logic ------------------

    def generate_document(self):
        title = self.doc_title_entry.get().strip()
        content = self.content_textbox.get("1.0", "end-1c").strip()
        suggested_filename = self.filename_entry.get().strip()

        if not title and not content:
            tkinter.messagebox.showerror("Error", "Please provide a Title or Content for the document.")
            return

        if not suggested_filename:
             suggested_filename = f"Generated_Doc_{datetime.now().strftime('%Y%m%d%H%M%S')}"

        if not suggested_filename.endswith(".docx"):
            suggested_filename += ".docx"

        # Ask the user where to save the file
        save_path = tkinter.filedialog.asksaveasfilename(
            title="Save Word Document",
            initialfile=suggested_filename,
            initialdir=self.last_saved_dir,
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")]
        )
        
        if not save_path:
            # User cancelled the save dialog
            return

        try:
            doc = Document()
            
            if title:
                doc.add_heading(title, 0)
                
            if content:
                # Add paragraphs handling simple new lines
                paragraphs = content.split('\n')
                for p in paragraphs:
                    if p.strip(): # Don't add empty stray paragraphs
                        doc.add_paragraph(p)

            # Save the file to the chosen path
            doc.save(save_path)
            
            # Remember the directory for next time
            self.last_saved_dir = os.path.dirname(save_path)
            
            tkinter.messagebox.showinfo("Success", f"Document successfully generated!\nSaved to:\n{save_path}")
            
        except Exception as e:
            tkinter.messagebox.showerror("Generation Error", f"An error occurred while generating the document:\n{e}")

    def open_folder(self):
        import platform
        import subprocess
        # Open the directory where the user last saved a file (or Documents by default)
        if platform.system() == "Windows":
            os.startfile(self.last_saved_dir)
        elif platform.system() == "Darwin":
            subprocess.Popen(["open", self.last_saved_dir])
        else:
            subprocess.Popen(["xdg-open", self.last_saved_dir])

if __name__ == "__main__":
    app = App()
    app.mainloop()
