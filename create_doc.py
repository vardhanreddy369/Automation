from docx import Document

def create_word_doc():
    doc = Document()
    doc.add_heading('Hello from Antigravity!', 0)
    
    p = doc.add_paragraph('You requested a new Word document. Here it is! I have typed this text directly into it programmatically.')
    p.add_run(' I hope this is what you were looking for!').bold = True
    
    doc.save('my_new_word_document.docx')

if __name__ == "__main__":
    create_word_doc()
